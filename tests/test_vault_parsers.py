"""Vault parsers: md/txt/html/docx/pdf + corrupt-file behavior."""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from vault import parsers


def test_md_txt_html(tmp_path):
    md = tmp_path / "a.md"
    md.write_text("# Title\n\nHello world", encoding="utf-8")
    text, pages, kind = parsers.parse_file(md)
    assert kind == "md" and "Hello world" in text and pages == 1
    tx = tmp_path / "b.txt"
    tx.write_text("plain words", encoding="utf-8")
    text, _, kind = parsers.parse_file(tx)
    assert kind == "txt" and "plain words" in text
    hx = tmp_path / "c.html"
    hx.write_text("<html><body><h1>Hi</h1><script>var x=1;</script><p>Body text</p></body></html>",
                  encoding="utf-8")
    text, _, kind = parsers.parse_file(hx)
    assert kind == "html" and "Body text" in text and "var x" not in text


def test_docx_headings_and_table(tmp_path):
    docx = pytest_importorskip_docx(tmp_path)
    text, pages, kind = parsers.parse_file(docx)
    assert kind == "docx" and "Pets" in text and "fluffy" in text


def pytest_importorskip_docx(tmp_path):
    try:
        from docx import Document
    except Exception:
        import pytest
        pytest.skip("python-docx not installed")
    p = tmp_path / "lease.docx"
    doc = Document()
    doc.add_heading("Pet Policy", level=1)
    doc.add_paragraph("Pets like fluffy cats are allowed with a deposit.")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "item"
    t.cell(0, 1).text = "fee"
    t.cell(1, 0).text = "dog"
    t.cell(1, 1).text = "200"
    doc.save(str(p))
    return p


def test_pdf_headings(tmp_path):
    try:
        import fitz
    except Exception:
        import pytest
        pytest.skip("pymupdf not installed")
    p = tmp_path / "lease.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Pet Policy\nPets are allowed with deposit of 200.")
    doc.save(str(p))
    doc.close()
    text, pages, kind = parsers.parse_file(p)
    assert kind == "pdf" and pages == 1 and "Pets are allowed" in text


def test_corrupt_file_raises_for_failed_status(tmp_path):
    bad = tmp_path / "bad.pdf"
    bad.write_bytes(b"%PDF-1.4 garbage \x00\xff not a real pdf trailer qqq")
    try:
        parsers.parse_file(bad)
        raised = False
    except Exception:
        raised = True
    assert raised is True
