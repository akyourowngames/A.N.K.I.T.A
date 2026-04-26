from __future__ import annotations

import json
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from pypdf import PdfReader

from jakata_agent.agent import JakataAgent
from jakata_agent.plan_validator import PlanValidator
from jakata_agent.router import IntentRouter
from jakata_agent.tasks.store import TaskStore
from jakata_agent.tools.base import Tool, ToolResult
from jakata_agent.tools.document import DocumentTool, register_document_tool
from jakata_agent.tools.registry import ToolRegistry


DOCUMENT_JSON = json.dumps(
    {
        "title": "AI Study Plan",
        "subtitle": "A practical weekly plan",
        "summary": "This plan helps a student study consistently with measurable checkpoints.",
        "sections": [
            {
                "heading": "Weekly Focus",
                "paragraphs": ["Use short daily sessions and reserve one review block for weak topics."],
                "bullets": ["Study math for 45 minutes.", "Review notes before sleeping."],
                "table": [["Day", "Focus"], ["Monday", "Math"], ["Tuesday", "Physics"]],
            },
            {
                "heading": "Next Steps",
                "paragraphs": ["Track marks and adjust the plan every Sunday."],
                "bullets": ["Keep the document updated."],
            },
        ],
        "sources": [{"title": "Study source", "url": "https://example.com/study", "note": "Supports spaced review."}],
    }
)


class FakeDocClient:
    def __init__(self, draft: str = DOCUMENT_JSON) -> None:
        self.draft = draft
        self.calls: list[tuple[str, str]] = []

    def complete_text(self, system_prompt: str, user_prompt: str, temperature: float = 0.7):
        del temperature
        self.calls.append((system_prompt, user_prompt))
        if "Available tools:" in user_prompt:
            return "router", '{"steps":[{"tool":"document","args":{"action":"create","prompt":"make a study plan","format":"both"},"reason":"create document"}]}'
        return "draft-model", self.draft

    def complete(self, messages, temperature: float = 0.7):
        del messages, temperature
        return "chat", "ok"

    def stream_complete(self, messages, temperature: float = 0.7):
        del messages, temperature
        yield "chat", "ok"


class FakeSearchTool(Tool):
    name = "search_web"
    description = "fake search"
    input_schema = {"type": "object", "properties": {}, "required": []}

    def run(self, args: dict) -> ToolResult:
        assert "query" in args
        return ToolResult(
            ok=True,
            summary="found",
            data={
                "results": [
                    {
                        "title": "Coffee market report",
                        "url": "https://example.com/coffee",
                        "content": "Specialty coffee demand is growing.",
                    }
                ]
            },
        )


class DummyMemory:
    def __init__(self) -> None:
        self.messages = []
        self.embedder = None
        self.store = SimpleNamespace(recent=lambda limit=10: [])
        self.knowledge_chunks = []

    def bootstrap_system_note(self):
        return "Memory empty."

    def persist_turn(self, messages):
        self.messages = list(messages)

    def learn_from_user_message(self, user_message: str):
        del user_message

    def load_session_messages(self):
        return []

    def retrieve(self, query: str):
        del query
        return SimpleNamespace(to_system_context=lambda: "", permanent_memories=[], knowledge_chunks=[], archived_chat_chunks=[])

    def graph_search(self, query: str):
        del query
        return []


def make_tool(tmp_path: Path, client: FakeDocClient | None = None, search_tool: Tool | None = None) -> DocumentTool:
    return DocumentTool(
        client=client or FakeDocClient(),
        output_dir=tmp_path / "generated" / "documents",
        template_dir=tmp_path / "templates",
        workspace_dir=tmp_path,
        search_tool=search_tool,
    )


def test_document_tool_creates_docx_and_pdf(tmp_path: Path):
    tool = make_tool(tmp_path)

    result = tool.run({"action": "create", "prompt": "make a study plan", "format": "both"})

    assert result.ok
    paths = [Path(path) for path in result.data["paths"]]
    assert {path.suffix for path in paths} == {".docx", ".pdf"}
    docx_path = next(path for path in paths if path.suffix == ".docx")
    pdf_path = next(path for path in paths if path.suffix == ".pdf")
    assert "AI Study Plan" in "\n".join(p.text for p in Document(str(docx_path)).paragraphs)
    assert len(PdfReader(str(pdf_path)).pages) >= 1


def test_document_tool_research_create_adds_sources(tmp_path: Path):
    tool = make_tool(tmp_path, search_tool=FakeSearchTool())

    result = tool.run({"action": "research_create", "prompt": "coffee shop market", "format": "docx"})

    assert result.ok
    assert result.data["source_count"] >= 1
    assert "https://example.com/coffee" in json.dumps(result.data["sources"])


def test_document_tool_edits_docx_with_replacements_and_append(tmp_path: Path):
    source = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("Hello {{name}}")
    doc.save(str(source))
    tool = make_tool(tmp_path)

    result = tool.run(
        {
            "action": "edit_docx",
            "source_path": str(source),
            "replacements": {"{{name}}": "Krish"},
            "append_text": "Final note added.",
        }
    )

    assert result.ok
    edited = Document(result.data["paths"][0])
    text = "\n".join(paragraph.text for paragraph in edited.paragraphs)
    assert "Hello Krish" in text
    assert "Final note added." in text


def test_document_tool_merge_and_split_pdf(tmp_path: Path):
    tool = make_tool(tmp_path)
    first = Path(tool.run({"action": "create", "content": "First PDF", "title": "First", "format": "pdf"}).data["paths"][0])
    second = Path(tool.run({"action": "create", "content": "Second PDF", "title": "Second", "format": "pdf"}).data["paths"][0])

    merged = tool.run({"action": "merge_pdf", "paths": [str(first), str(second)]})
    assert merged.ok
    merged_path = Path(merged.data["paths"][0])
    assert len(PdfReader(str(merged_path)).pages) >= 2

    split = tool.run({"action": "split_pdf", "source_path": str(merged_path), "page_ranges": "1"})
    assert split.ok
    assert len(PdfReader(split.data["paths"][0]).pages) == 1


def test_document_tool_extracts_pdf_text(tmp_path: Path):
    tool = make_tool(tmp_path, client=FakeDocClient(draft="not json"))
    pdf = Path(tool.run({"action": "create", "content": "Extract this sentence.", "title": "Extract", "format": "pdf"}).data["paths"][0])

    result = tool.run({"action": "extract", "source_path": str(pdf)})

    assert result.ok
    assert "Extract" in result.data["preview"] or "Extract this sentence" in result.data["preview"]


def test_agent_routes_document_creation_through_real_tool(tmp_path: Path):
    client = FakeDocClient()
    tools = ToolRegistry()
    register_document_tool(
        tools,
        client=client,
        output_dir=tmp_path / "data" / "generated" / "documents",
        template_dir=tmp_path / "data" / "document_templates",
        workspace_dir=tmp_path,
    )
    agent = JakataAgent(
        settings=SimpleNamespace(
            session_id="s1",
            data_dir=tmp_path / "data",
            workspace_dir=tmp_path,
            router_tool_limit=0,
            router_min_tool_score=0,
            approval_policy="auto_safe",
        ),
        client=client,
        tools=tools,
        memory=DummyMemory(),
        router=IntentRouter(client),
        validator=PlanValidator(),
        task_store=TaskStore(tmp_path / "jakata.db"),
    )

    model, content = agent.reply("create a pdf and docx study plan")

    assert model == "local:tool"
    assert "Created DOCX and PDF documents" in content
    assert list((tmp_path / "data" / "generated" / "documents").glob("*.docx"))
    assert list((tmp_path / "data" / "generated" / "documents").glob("*.pdf"))
