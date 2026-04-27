from __future__ import annotations

import json
import re
import shutil
import subprocess
import time
from html import escape
from pathlib import Path
from typing import Any

from jakata_agent.llm import TextCompletionClient
from jakata_agent.prompts import load_prompt
from jakata_agent.tools.base import Tool, ToolResult
from jakata_agent.tools.opening import open_target
from jakata_agent.tools.registry import ToolRegistry


DOCUMENT_DRAFTER_PROMPT = load_prompt("document/drafter.md")
_SUPPORTED_CREATE_FORMATS = {"docx", "pdf", "both"}
_PDF_RANGE_RE = re.compile(r"^\d+(-\d+)?$")


class DocumentTool(Tool):
    name = "document"
    description = (
        "Create, research, edit, convert, extract, merge, split, and annotate polished document files only when "
        "the user wants a saved/exported DOCX/PDF/TXT artifact or asks to modify an existing document. "
        "Do not use for ordinary in-chat planning, web page copy, outlines, ideas, or answers that the user wants "
        "shown in chat instead of saved to a file. Supports DOCX and PDF outputs, reusable DOCX templates, "
        "source-backed research documents, PDF page operations, and returns local file paths under the generated documents folder."
    )
    safety = "write"
    categories = ("documents", "productivity", "artifact_creation")
    aliases = ("docx", "pdf", "report", "save as document", "make a file")
    use_with = ("search_web", "camera", "ocr", "screen", "memory", "weather")
    daily_uses = (
        "Turn research, camera observations, OCR text, notes, plans, or weather context into saved DOCX/PDF/TXT artifacts.",
        "Edit, extract, convert, merge, split, annotate, and template documents.",
    )
    grounding = "Creates or modifies local files under the generated documents folder and returns saved paths plus preview/source metadata."
    output_capabilities = ("docx", "pdf", "txt", "paths", "preview", "sources")
    input_schema = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "enum": [
                    "create",
                    "research_create",
                    "edit_docx",
                    "fill_template",
                    "convert",
                    "extract",
                    "merge_pdf",
                    "split_pdf",
                    "annotate_pdf",
                ],
                "description": "Document operation to run.",
            },
            "prompt": {"type": "string", "description": "Goal or instruction for generated or edited content."},
            "content": {"type": "string", "description": "Source text to turn into a document."},
            "title": {"type": "string", "description": "Optional document title."},
            "format": {"type": "string", "enum": ["docx", "pdf", "both", "txt"], "description": "Output format."},
            "output_name": {"type": "string", "description": "Optional output filename or stem."},
            "source_path": {"type": "string", "description": "Existing document path for edit, convert, extract, split, or annotate."},
            "paths": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Input PDF paths for merge_pdf.",
            },
            "page_ranges": {
                "type": "string",
                "description": "1-based PDF page ranges like 1-3,5,8-10 for split_pdf.",
            },
            "replacements": {
                "type": "object",
                "description": "Text replacements for edit_docx or template variables for fill_template.",
            },
            "append_text": {"type": "string", "description": "Text to append during edit_docx."},
            "watermark": {"type": "string", "description": "Watermark text for annotate_pdf."},
            "annotation": {"type": "string", "description": "Visible note text for annotate_pdf."},
            "research_query": {"type": "string", "description": "Search query for research_create. Defaults to prompt."},
            "open_after": {"type": "boolean", "description": "Whether the caller should open the output after creation."},
        },
        "required": ["action"],
        "additionalProperties": False,
    }

    def __init__(
        self,
        *,
        client: TextCompletionClient,
        output_dir: str | Path,
        template_dir: str | Path,
        workspace_dir: str | Path,
        search_tool: Tool | None = None,
    ) -> None:
        self.client = client
        self.output_dir = Path(output_dir).expanduser().resolve()
        self.template_dir = Path(template_dir).expanduser().resolve()
        self.workspace_dir = Path(workspace_dir).expanduser().resolve()
        self.search_tool = search_tool

    def normalize_args(self, args: dict[str, Any]) -> dict[str, Any]:
        args.setdefault("action", "create")
        args.setdefault("format", "both")
        args.setdefault("open_after", False)
        return args

    def run(self, args: dict[str, Any]) -> ToolResult:
        args = self.normalize_args(dict(args))
        action = str(args.get("action", "create")).strip().lower()
        try:
            if action in {"create", "research_create"}:
                return self._create(args, research=action == "research_create")
            if action == "edit_docx":
                return self._edit_docx(args)
            if action == "fill_template":
                return self._fill_template(args)
            if action == "convert":
                return self._convert(args)
            if action == "extract":
                return self._extract(args)
            if action == "merge_pdf":
                return self._merge_pdf(args)
            if action == "split_pdf":
                return self._split_pdf(args)
            if action == "annotate_pdf":
                return self._annotate_pdf(args)
        except MissingDocumentDependency as exc:
            return ToolResult(ok=False, summary=str(exc), data={}, error="missing_dependency")
        except Exception as exc:  # noqa: BLE001
            return ToolResult(ok=False, summary=f"Document action failed: {exc}", data={}, error="document_failed")
        return ToolResult(ok=False, summary=f"Unsupported document action: {action}", data={}, error="unsupported_action")

    def render(self, data: dict[str, Any]) -> str:
        summary = str(data.get("summary", "")).strip() or "Document action complete."
        lines = [f"{summary}, sir."]
        paths = data.get("paths") or []
        if isinstance(paths, list) and paths:
            lines.append("")
            lines.append("The file is saved here:" if len(paths) == 1 else "The files are saved here:")
            for path in paths:
                lines.append(f"- {path}")
        source_count = int(data.get("source_count", 0) or 0)
        if source_count:
            lines.append(f"\nSources used: {source_count}")
        preview = str(data.get("preview", "")).strip()
        if preview:
            lines.append(f"\nPreview:\n{preview[:1200]}")
        return "\n".join(lines).strip()

    def _create(self, args: dict[str, Any], *, research: bool) -> ToolResult:
        output_format = self._create_format(args.get("format", "both"))
        prompt = str(args.get("prompt", "")).strip()
        content = str(args.get("content", "")).strip()
        title = str(args.get("title", "")).strip()
        if not prompt and not content and not title:
            return ToolResult(ok=False, summary="Document create requires prompt, content, or title.", data={}, error="missing_prompt")

        sources = self._research_sources(args, prompt=prompt, enabled=research)
        spec = self._draft_spec(prompt=prompt, content=content, title=title, sources=sources)
        paths = self._write_spec_outputs(spec, output_format=output_format, output_name=str(args.get("output_name", "")).strip())
        open_results = self._open_outputs(paths) if bool(args.get("open_after", False)) else []
        summary = f"Created {output_format.upper()} document" if output_format != "both" else "Created DOCX and PDF documents"
        if research:
            summary = "Created researched DOCX/PDF document" if output_format == "both" else f"Created researched {output_format.upper()} document"
        return ToolResult(
            ok=True,
            summary=f"The document is ready, sir. {summary}: {', '.join(path.name for path in paths)}",
            data={
                "summary": summary,
                "paths": [str(path) for path in paths],
                "format": output_format,
                "title": spec.get("title", ""),
                "source_count": len(spec.get("sources", [])),
                "sources": spec.get("sources", []),
                "preview": self._spec_preview(spec),
                "opened": bool(open_results) and all(item.get("opened") for item in open_results),
                "open_results": open_results,
            },
        )

    def _edit_docx(self, args: dict[str, Any]) -> ToolResult:
        Document = self._docx_document()
        source = self._resolve_existing_path(str(args.get("source_path", "")).strip())
        if source.suffix.lower() != ".docx":
            return ToolResult(ok=False, summary="edit_docx requires a .docx source_path.", data={}, error="not_docx")
        doc = Document(str(source))
        replacements = self._dict_arg(args.get("replacements"))
        changed = 0
        for paragraph in doc.paragraphs:
            changed += self._replace_paragraph_text(paragraph, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        changed += self._replace_paragraph_text(paragraph, replacements)

        append_text = str(args.get("append_text", "")).strip()
        prompt = str(args.get("prompt", "")).strip()
        if not append_text and prompt:
            spec = self._draft_spec(prompt=prompt, content="", title="", sources=[])
            append_text = self._spec_to_plain_text(spec)
        if append_text:
            doc.add_page_break()
            doc.add_heading("Added Notes", level=1)
            for paragraph in self._paragraphs_from_text(append_text):
                doc.add_paragraph(paragraph)

        output = self._target_path(args.get("output_name") or f"{source.stem}-edited", ".docx")
        doc.save(str(output))
        summary = f"Edited DOCX document: {output.name}"
        return ToolResult(
            ok=True,
            summary=summary,
            data={"summary": summary, "paths": [str(output)], "replacements": changed, "preview": append_text[:800]},
        )

    def _fill_template(self, args: dict[str, Any]) -> ToolResult:
        try:
            from docxtpl import DocxTemplate
        except ImportError as exc:
            raise MissingDocumentDependency("docxtpl is not installed. Run `pip install -r requirements.txt`.") from exc
        source = self._resolve_existing_path(str(args.get("source_path", "")).strip(), template_ok=True)
        if source.suffix.lower() != ".docx":
            return ToolResult(ok=False, summary="fill_template requires a .docx template source_path.", data={}, error="not_docx")
        context = self._dict_arg(args.get("replacements"))
        if not context:
            return ToolResult(ok=False, summary="fill_template requires replacements/template context.", data={}, error="missing_context")
        doc = DocxTemplate(str(source))
        doc.render(context)
        output = self._target_path(args.get("output_name") or f"{source.stem}-filled", ".docx")
        doc.save(str(output))
        summary = f"Filled DOCX template: {output.name}"
        return ToolResult(ok=True, summary=summary, data={"summary": summary, "paths": [str(output)], "keys": sorted(context.keys())})

    def _convert(self, args: dict[str, Any]) -> ToolResult:
        source = self._resolve_existing_path(str(args.get("source_path", "")).strip())
        target_format = str(args.get("format", "pdf")).strip().lower() or "pdf"
        if target_format not in {"pdf", "docx", "txt"}:
            return ToolResult(ok=False, summary="convert format must be pdf, docx, or txt.", data={}, error="unsupported_format")
        if source.suffix.lower() == f".{target_format}":
            output = self._target_path(args.get("output_name") or f"{source.stem}-copy", source.suffix)
            shutil.copyfile(source, output)
        elif target_format == "pdf":
            output = self._target_path(args.get("output_name") or source.stem, ".pdf")
            self._convert_to_pdf(source, output)
        elif target_format == "docx":
            output = self._target_path(args.get("output_name") or source.stem, ".docx")
            self._text_to_docx(self._extract_text(source), output, title=source.stem)
        else:
            output = self._target_path(args.get("output_name") or source.stem, ".txt")
            output.write_text(self._extract_text(source), encoding="utf-8")
        summary = f"Converted document to {target_format.upper()}: {output.name}"
        return ToolResult(ok=True, summary=summary, data={"summary": summary, "paths": [str(output)], "source_path": str(source)})

    def _extract(self, args: dict[str, Any]) -> ToolResult:
        source = self._resolve_existing_path(str(args.get("source_path", "")).strip())
        text = self._extract_text(source)
        output = self._target_path(args.get("output_name") or f"{source.stem}-extracted", ".txt")
        output.write_text(text, encoding="utf-8")
        summary = f"Extracted text from document: {source.name}"
        return ToolResult(
            ok=True,
            summary=summary,
            data={"summary": summary, "paths": [str(output)], "source_path": str(source), "characters": len(text), "preview": text[:1200]},
        )

    def _merge_pdf(self, args: dict[str, Any]) -> ToolResult:
        PdfWriter, _ = self._pypdf()
        raw_paths = args.get("paths") or []
        if not isinstance(raw_paths, list) or len(raw_paths) < 2:
            return ToolResult(ok=False, summary="merge_pdf requires at least two PDF paths.", data={}, error="missing_paths")
        writer = PdfWriter()
        inputs = [self._resolve_existing_path(str(path).strip()) for path in raw_paths]
        for path in inputs:
            if path.suffix.lower() != ".pdf":
                return ToolResult(ok=False, summary=f"merge_pdf only accepts PDFs: {path}", data={}, error="not_pdf")
            writer.append(str(path))
        output = self._target_path(args.get("output_name") or "merged-document", ".pdf")
        with output.open("wb") as handle:
            writer.write(handle)
        summary = f"Merged {len(inputs)} PDFs: {output.name}"
        return ToolResult(ok=True, summary=summary, data={"summary": summary, "paths": [str(output)], "inputs": [str(path) for path in inputs]})

    def _split_pdf(self, args: dict[str, Any]) -> ToolResult:
        PdfWriter, PdfReader = self._pypdf()
        source = self._resolve_existing_path(str(args.get("source_path", "")).strip())
        if source.suffix.lower() != ".pdf":
            return ToolResult(ok=False, summary="split_pdf requires a PDF source_path.", data={}, error="not_pdf")
        reader = PdfReader(str(source))
        pages = self._parse_page_ranges(str(args.get("page_ranges", "")).strip(), len(reader.pages))
        writer = PdfWriter()
        for page_index in pages:
            writer.add_page(reader.pages[page_index])
        output = self._target_path(args.get("output_name") or f"{source.stem}-pages", ".pdf")
        with output.open("wb") as handle:
            writer.write(handle)
        summary = f"Split PDF pages {self._range_summary(pages)}: {output.name}"
        return ToolResult(ok=True, summary=summary, data={"summary": summary, "paths": [str(output)], "pages": [page + 1 for page in pages]})

    def _annotate_pdf(self, args: dict[str, Any]) -> ToolResult:
        pymupdf = self._pymupdf()
        source = self._resolve_existing_path(str(args.get("source_path", "")).strip())
        if source.suffix.lower() != ".pdf":
            return ToolResult(ok=False, summary="annotate_pdf requires a PDF source_path.", data={}, error="not_pdf")
        watermark = str(args.get("watermark", "")).strip()
        annotation = str(args.get("annotation", "")).strip()
        if not watermark and not annotation:
            return ToolResult(ok=False, summary="annotate_pdf requires watermark or annotation text.", data={}, error="missing_annotation")
        doc = pymupdf.open(str(source))
        for page in doc:
            rect = page.rect
            if watermark:
                page.insert_text(
                    (rect.width * 0.18, rect.height * 0.52),
                    watermark,
                    fontsize=42,
                    color=(0.78, 0.78, 0.78),
                    rotate=30,
                    overlay=True,
                )
            if annotation and page.number == 0:
                note_rect = pymupdf.Rect(48, 36, rect.width - 48, 96)
                page.insert_textbox(note_rect, annotation, fontsize=10, color=(0.1, 0.1, 0.1), overlay=True)
        output = self._target_path(args.get("output_name") or f"{source.stem}-annotated", ".pdf")
        doc.save(str(output))
        doc.close()
        summary = f"Annotated PDF: {output.name}"
        return ToolResult(ok=True, summary=summary, data={"summary": summary, "paths": [str(output)], "source_path": str(source)})

    def _research_sources(self, args: dict[str, Any], *, prompt: str, enabled: bool) -> list[dict[str, str]]:
        if not enabled or self.search_tool is None:
            return []
        query = str(args.get("research_query", "")).strip() or prompt
        if not query:
            return []
        result = self.search_tool.run({"query": query, "max_results": 5, "topic": "general"})
        if not result.ok:
            return []
        sources: list[dict[str, str]] = []
        for item in result.data.get("results", []) if isinstance(result.data, dict) else []:
            if not isinstance(item, dict):
                continue
            sources.append(
                {
                    "title": str(item.get("title", "")).strip(),
                    "url": str(item.get("url", "")).strip(),
                    "note": str(item.get("content", "")).strip(),
                }
            )
        return sources

    def _draft_spec(self, *, prompt: str, content: str, title: str, sources: list[dict[str, str]]) -> dict[str, Any]:
        payload = {
            "title": title,
            "prompt": prompt,
            "content": content,
            "research_sources": sources,
        }
        if not self.client:
            return self._fallback_spec(title=title, prompt=prompt, content=content, sources=sources)
        try:
            _, raw = self.client.complete_text(DOCUMENT_DRAFTER_PROMPT, json.dumps(payload, ensure_ascii=False), temperature=0.35)
            parsed = self._extract_json(raw)
            return self._normalize_spec(parsed, fallback_title=title or prompt, sources=sources)
        except Exception:
            return self._fallback_spec(title=title, prompt=prompt, content=content, sources=sources)

    def _write_spec_outputs(self, spec: dict[str, Any], *, output_format: str, output_name: str) -> list[Path]:
        paths: list[Path] = []
        stem = output_name or spec.get("title") or "document"
        if output_format in {"docx", "both"}:
            docx_path = self._target_path(stem, ".docx")
            self._write_docx(spec, docx_path)
            paths.append(docx_path)
        if output_format in {"pdf", "both"}:
            pdf_path = self._target_path(stem, ".pdf")
            self._write_pdf(spec, pdf_path)
            paths.append(pdf_path)
        return paths

    @staticmethod
    def _open_outputs(paths: list[Path]) -> list[dict[str, Any]]:
        results: list[dict[str, Any]] = []
        for path in paths:
            opened = open_target(path, wait_seconds=1.5)
            results.append(opened.as_dict())
        return results

    def _write_docx(self, spec: dict[str, Any], path: Path) -> None:
        Document = self._docx_document()
        doc = Document()
        doc.core_properties.title = str(spec.get("title", "Document"))
        doc.add_heading(str(spec.get("title", "Document")), level=0)
        subtitle = str(spec.get("subtitle", "")).strip()
        if subtitle:
            doc.add_paragraph(subtitle)
        summary = str(spec.get("summary", "")).strip()
        if summary:
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(summary)
        for section in spec.get("sections", []):
            if not isinstance(section, dict):
                continue
            heading = str(section.get("heading", "")).strip()
            if heading:
                doc.add_heading(heading, level=1)
            for paragraph in section.get("paragraphs", []) or []:
                text = str(paragraph).strip()
                if text:
                    doc.add_paragraph(text)
            for bullet in section.get("bullets", []) or []:
                text = str(bullet).strip()
                if text:
                    doc.add_paragraph(text, style="List Bullet")
            table = section.get("table")
            if self._valid_table(table):
                rows = [[str(cell) for cell in row] for row in table]
                doc_table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
                doc_table.style = "Table Grid"
                for row_index, row in enumerate(rows):
                    for col_index, value in enumerate(row):
                        doc_table.cell(row_index, col_index).text = value
        sources = spec.get("sources", [])
        if sources:
            doc.add_heading("Sources", level=1)
            for source in sources:
                doc.add_paragraph(self._source_line(source), style="List Number")
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))

    def _write_pdf(self, spec: dict[str, Any], path: Path) -> None:
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except ImportError as exc:
            raise MissingDocumentDependency("reportlab is not installed. Run `pip install -r requirements.txt`.") from exc

        path.parent.mkdir(parents=True, exist_ok=True)
        styles = getSampleStyleSheet()
        title_style = styles["Title"]
        heading_style = styles["Heading1"]
        body_style = styles["BodyText"]
        body_style.spaceAfter = 8
        bullet_style = ParagraphStyle("DocumentBullet", parent=body_style, leftIndent=18, bulletIndent=8)
        story: list[Any] = [Paragraph(escape(str(spec.get("title", "Document"))), title_style)]
        subtitle = str(spec.get("subtitle", "")).strip()
        if subtitle:
            story.extend([Spacer(1, 0.08 * inch), Paragraph(escape(subtitle), styles["Italic"])])
        summary = str(spec.get("summary", "")).strip()
        if summary:
            story.extend([Spacer(1, 0.18 * inch), Paragraph("Executive Summary", heading_style), Paragraph(escape(summary), body_style)])
        for section in spec.get("sections", []):
            if not isinstance(section, dict):
                continue
            heading = str(section.get("heading", "")).strip()
            if heading:
                story.extend([Spacer(1, 0.14 * inch), Paragraph(escape(heading), heading_style)])
            for paragraph in section.get("paragraphs", []) or []:
                text = str(paragraph).strip()
                if text:
                    story.append(Paragraph(escape(text), body_style))
            for bullet in section.get("bullets", []) or []:
                text = str(bullet).strip()
                if text:
                    story.append(Paragraph(escape(text), bullet_style, bulletText="•"))
            table = section.get("table")
            if self._valid_table(table):
                rows = [[Paragraph(escape(str(cell)), body_style) for cell in row] for row in table]
                pdf_table = Table(rows, repeatRows=1, hAlign="LEFT")
                pdf_table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDEFF3")),
                            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B9C0CC")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 6),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ]
                    )
                )
                story.extend([Spacer(1, 0.08 * inch), pdf_table, Spacer(1, 0.08 * inch)])
        sources = spec.get("sources", [])
        if sources:
            story.extend([Spacer(1, 0.16 * inch), Paragraph("Sources", heading_style)])
            for index, source in enumerate(sources, 1):
                story.append(Paragraph(escape(f"{index}. {self._source_line(source)}"), body_style))
        doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=48, leftMargin=48, topMargin=52, bottomMargin=52)
        doc.build(story)

    def _convert_to_pdf(self, source: Path, output: Path) -> None:
        if source.suffix.lower() == ".docx" and self._try_libreoffice_pdf(source, output):
            return
        text = self._extract_text(source)
        spec = self._fallback_spec(title=source.stem, prompt="", content=text, sources=[])
        self._write_pdf(spec, output)

    def _try_libreoffice_pdf(self, source: Path, output: Path) -> bool:
        soffice = self._find_soffice()
        if not soffice:
            return False
        output.parent.mkdir(parents=True, exist_ok=True)
        temp_dir = output.parent / f".lo-{int(time.time() * 1000)}"
        temp_dir.mkdir(parents=True, exist_ok=True)
        try:
            completed = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(temp_dir), str(source)],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=90,
            )
            generated = temp_dir / f"{source.stem}.pdf"
            if completed.returncode == 0 and generated.exists():
                shutil.move(str(generated), str(output))
                return True
            return False
        except Exception:
            return False
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    @staticmethod
    def _find_soffice() -> str:
        for candidate in [
            shutil.which("soffice"),
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]:
            if candidate and Path(candidate).exists():
                return str(candidate)
        return ""

    def _text_to_docx(self, text: str, output: Path, *, title: str) -> None:
        spec = self._fallback_spec(title=title, prompt="", content=text, sources=[])
        self._write_docx(spec, output)

    def _extract_text(self, source: Path) -> str:
        suffix = source.suffix.lower()
        if suffix == ".docx":
            Document = self._docx_document()
            doc = Document(str(source))
            lines = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    lines.append(" | ".join(cell.text.strip() for cell in row.cells))
            return "\n".join(lines)
        if suffix == ".pdf":
            pymupdf = self._pymupdf()
            doc = pymupdf.open(str(source))
            try:
                return "\n".join(page.get_text("text").strip() for page in doc if page.get_text("text").strip())
            finally:
                doc.close()
        return source.read_text(encoding="utf-8", errors="replace")

    def _resolve_existing_path(self, raw_path: str, *, template_ok: bool = False) -> Path:
        if not raw_path:
            raise ValueError("source_path is required.")
        candidates = [Path(raw_path).expanduser()]
        if not candidates[0].is_absolute():
            candidates.append(self.workspace_dir / raw_path)
            candidates.append(self.output_dir / raw_path)
            if template_ok:
                candidates.append(self.template_dir / raw_path)
        for candidate in candidates:
            resolved = candidate.resolve()
            if resolved.exists() and resolved.is_file():
                return resolved
        raise FileNotFoundError(raw_path)

    def _target_path(self, raw_name: Any, suffix: str) -> Path:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        name = str(raw_name or "document").strip()
        name = Path(name).name
        stem = Path(name).stem if Path(name).suffix else name
        stem = self._slug(stem) or "document"
        path = self.output_dir / f"{stem}{suffix}"
        if not path.exists():
            return path
        for index in range(2, 1000):
            candidate = self.output_dir / f"{stem}-{index}{suffix}"
            if not candidate.exists():
                return candidate
        return self.output_dir / f"{stem}-{int(time.time())}{suffix}"

    @staticmethod
    def _create_format(value: Any) -> str:
        output_format = str(value or "both").strip().lower()
        if output_format not in _SUPPORTED_CREATE_FORMATS:
            return "both"
        return output_format

    @staticmethod
    def _slug(value: str) -> str:
        value = re.sub(r"[^A-Za-z0-9._ -]+", "", value).strip().replace(" ", "-")
        value = re.sub(r"-{2,}", "-", value)
        return value[:90].strip(".-")

    @staticmethod
    def _dict_arg(value: Any) -> dict[str, Any]:
        if isinstance(value, dict):
            return value
        if isinstance(value, str) and value.strip():
            try:
                parsed = json.loads(value)
                return parsed if isinstance(parsed, dict) else {}
            except json.JSONDecodeError:
                return {}
        return {}

    @staticmethod
    def _extract_json(raw: str) -> dict[str, Any]:
        cleaned = raw.strip().replace("```json", "").replace("```", "").strip()
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start >= 0 and end >= start:
            cleaned = cleaned[start : end + 1]
        parsed = json.loads(cleaned)
        return parsed if isinstance(parsed, dict) else {}

    @staticmethod
    def _normalize_spec(spec: dict[str, Any], *, fallback_title: str, sources: list[dict[str, str]]) -> dict[str, Any]:
        title = str(spec.get("title") or fallback_title or "Document").strip()
        model_sources = spec.get("sources") if isinstance(spec.get("sources"), list) else []
        normalized_sources = DocumentTool._merge_sources(model_sources, sources)
        sections = spec.get("sections") if isinstance(spec.get("sections"), list) else []
        if not sections:
            sections = [{"heading": "Overview", "paragraphs": [str(spec.get("summary") or fallback_title or "").strip()], "bullets": []}]
        return {
            "title": title,
            "subtitle": str(spec.get("subtitle", "")).strip(),
            "summary": str(spec.get("summary", "")).strip(),
            "sections": sections,
            "sources": normalized_sources or sources,
        }

    @staticmethod
    def _merge_sources(model_sources: list[Any], research_sources: list[dict[str, str]]) -> list[Any]:
        merged: list[Any] = []
        seen: set[str] = set()
        for source in [*research_sources, *model_sources]:
            if not isinstance(source, dict):
                continue
            key = str(source.get("url") or source.get("title") or source).strip()
            if key and key not in seen:
                seen.add(key)
                merged.append(source)
        return merged

    def _fallback_spec(self, *, title: str, prompt: str, content: str, sources: list[dict[str, str]]) -> dict[str, Any]:
        doc_title = title or prompt or "Document"
        paragraphs = self._paragraphs_from_text(content or prompt or doc_title)
        if not paragraphs:
            paragraphs = [doc_title]
        if not content.strip() and prompt.strip():
            topic = self._clean_topic(prompt)
            doc_title = title or topic
            return {
                "title": doc_title[:120],
                "subtitle": "Practical, ready-to-use document",
                "summary": f"This document gives a clear, structured overview of {topic.lower()} with practical guidance, priorities, and next steps.",
                "sections": [
                    {
                        "heading": "Overview",
                        "paragraphs": [
                            f"{topic} benefits from a focused structure: define the goal, separate the key ideas, and turn each idea into specific actions.",
                            "The document below is organized for quick reading while still giving enough detail to act on immediately.",
                        ],
                        "bullets": [
                            "Start with the main outcome and keep every section tied to that outcome.",
                            "Prefer concrete examples, measurable checkpoints, and simple language.",
                            "Review and refine the document after new evidence or feedback is available.",
                        ],
                    },
                    {
                        "heading": "Recommended Plan",
                        "paragraphs": [
                            "Use the following plan as a working draft. Each step is intentionally practical so it can be converted into tasks, notes, or a final presentation.",
                        ],
                        "bullets": [
                            "Clarify the audience and purpose before adding detail.",
                            "Group related information under short headings.",
                            "Add supporting evidence, examples, or references where claims matter.",
                            "End with next steps so the document leads to action.",
                        ],
                        "table": [["Stage", "Output"], ["Plan", "Goal, audience, and outline"], ["Draft", "Structured sections with examples"], ["Review", "Gaps, risks, and improvements"]],
                    },
                    {
                        "heading": "Next Steps",
                        "paragraphs": [
                            "Improve this draft by adding project-specific data, names, dates, links, and any required formatting rules.",
                        ],
                        "bullets": [
                            "Verify facts before sharing externally.",
                            "Add source links or citations for research-heavy sections.",
                            "Export to PDF for sharing and keep DOCX as the editable source.",
                        ],
                    },
                ],
                "sources": sources,
            }
        return {
            "title": doc_title[:120],
            "subtitle": "",
            "summary": paragraphs[0],
            "sections": [{"heading": "Overview", "paragraphs": paragraphs, "bullets": []}],
            "sources": sources,
        }

    @staticmethod
    def _clean_topic(prompt: str) -> str:
        topic = re.sub(
            r"\b(create|make|write|generate|professional|short|detailed|docx|pdf|document|report|about|on|for|me|a|an|the)\b",
            " ",
            prompt,
            flags=re.IGNORECASE,
        )
        topic = " ".join(topic.split()).strip(" .,:;-")
        return topic.title() if topic else prompt.strip().title() or "Document"

    @staticmethod
    def _paragraphs_from_text(text: str) -> list[str]:
        blocks = [block.strip() for block in re.split(r"\n\s*\n", text.strip()) if block.strip()]
        if blocks:
            return blocks
        return [line.strip() for line in text.splitlines() if line.strip()]

    def _spec_to_plain_text(self, spec: dict[str, Any]) -> str:
        lines = [str(spec.get("title", "")).strip(), str(spec.get("summary", "")).strip()]
        for section in spec.get("sections", []):
            if not isinstance(section, dict):
                continue
            lines.append(str(section.get("heading", "")).strip())
            lines.extend(str(item).strip() for item in section.get("paragraphs", []) or [])
            lines.extend(f"- {str(item).strip()}" for item in section.get("bullets", []) or [])
        return "\n\n".join(line for line in lines if line)

    def _spec_preview(self, spec: dict[str, Any]) -> str:
        return self._spec_to_plain_text(spec)[:1200]

    @staticmethod
    def _source_line(source: Any) -> str:
        if not isinstance(source, dict):
            return str(source)
        title = str(source.get("title", "")).strip() or "Source"
        url = str(source.get("url", "")).strip()
        note = str(source.get("note", "")).strip()
        parts = [title]
        if url:
            parts.append(url)
        if note:
            parts.append(note[:240])
        return " - ".join(parts)

    @staticmethod
    def _valid_table(table: Any) -> bool:
        return isinstance(table, list) and len(table) >= 2 and all(isinstance(row, list) and row for row in table)

    @staticmethod
    def _replace_paragraph_text(paragraph: Any, replacements: dict[str, Any]) -> int:
        if not replacements:
            return 0
        original = paragraph.text
        updated = original
        for key, value in replacements.items():
            updated = updated.replace(str(key), str(value))
        if updated == original:
            return 0
        for run in list(paragraph.runs):
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = updated
        else:
            paragraph.add_run(updated)
        return 1

    @staticmethod
    def _parse_page_ranges(raw: str, total_pages: int) -> list[int]:
        if total_pages <= 0:
            return []
        if not raw:
            return list(range(total_pages))
        pages: list[int] = []
        for item in [part.strip() for part in raw.split(",") if part.strip()]:
            if not _PDF_RANGE_RE.match(item):
                raise ValueError(f"Invalid page range: {item}")
            if "-" in item:
                start, end = [int(part) for part in item.split("-", 1)]
            else:
                start = end = int(item)
            if start < 1 or end < start or end > total_pages:
                raise ValueError(f"Page range out of bounds: {item}")
            pages.extend(range(start - 1, end))
        return sorted(dict.fromkeys(pages))

    @staticmethod
    def _range_summary(pages: list[int]) -> str:
        return ",".join(str(page + 1) for page in pages)

    @staticmethod
    def _docx_document():
        try:
            from docx import Document
        except ImportError as exc:
            raise MissingDocumentDependency("python-docx is not installed. Run `pip install -r requirements.txt`.") from exc
        return Document

    @staticmethod
    def _pypdf():
        try:
            from pypdf import PdfReader, PdfWriter
        except ImportError as exc:
            raise MissingDocumentDependency("pypdf is not installed. Run `pip install -r requirements.txt`.") from exc
        return PdfWriter, PdfReader

    @staticmethod
    def _pymupdf():
        try:
            import pymupdf
        except ImportError as exc:
            raise MissingDocumentDependency("PyMuPDF is not installed. Run `pip install -r requirements.txt`.") from exc
        return pymupdf


class MissingDocumentDependency(RuntimeError):
    pass


def register_document_tool(
    registry: ToolRegistry,
    *,
    client: TextCompletionClient,
    output_dir: str | Path,
    template_dir: str | Path,
    workspace_dir: str | Path,
    search_tool: Tool | None = None,
) -> None:
    registry.register(
        DocumentTool(
            client=client,
            output_dir=output_dir,
            template_dir=template_dir,
            workspace_dir=workspace_dir,
            search_tool=search_tool,
        )
    )
