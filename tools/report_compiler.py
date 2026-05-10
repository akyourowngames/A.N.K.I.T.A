from __future__ import annotations

import json
import os
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from html import escape
from importlib.util import find_spec
from pathlib import Path
from typing import Any

from tools.registry import ToolInputError, optional_text


DEFAULT_CONFIG_PATH = Path("config/report_compiler.json")
SUPPORTED_FORMATS = {"md", "txt", "docx", "pdf"}


def compiler_status(_params: dict[str, Any] | None = None) -> dict[str, Any]:
    config = load_config()
    templates = load_templates(config)
    dependencies = dependency_status()
    return {
        "summary": f"{config.get('agent_name', 'Report Compiler')} ready.",
        "config_path": str(config_path()),
        "output_dir": str(resolve_path(str(config.get("output_dir")))),
        "templates_dir": str(resolve_path(str(config.get("templates_dir")))),
        "supported_formats": sorted(SUPPORTED_FORMATS),
        "default_format": config.get("default_format"),
        "default_template": config.get("default_template"),
        "templates": sorted(templates.keys()),
        "dependencies": dependencies,
    }


def compiler_list_templates(_params: dict[str, Any] | None = None) -> dict[str, Any]:
    config = load_config()
    templates = load_templates(config)
    return {
        "count": len(templates),
        "templates": [
            {
                "name": name,
                "description": template.get("description", ""),
                "sections": template.get("sections", []),
                "required_fields": template.get("required_fields", []),
                "citation_style": template.get("citation_style", config.get("citation_style")),
            }
            for name, template in sorted(templates.items())
        ],
    }


def compiler_validate(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    template_name = optional_text(params, "template", str(config.get("default_template") or "blank"))
    templates = load_templates(config)
    template = templates.get(template_name, {})
    content = params.get("content")
    normalized, errors, warnings = validate_content(content, template)
    return {
        "valid": not errors,
        "errors": errors,
        "warnings": warnings,
        "template": template_name,
        "normalized_content": normalized,
    }


def compiler_preview(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    template_name = optional_text(params, "template", str(config.get("default_template") or "blank"))
    templates = load_templates(config)
    normalized, errors, warnings = validate_content(params.get("content"), templates.get(template_name, {}))
    if errors:
        raise ToolInputError("; ".join(errors))
    preview = render_to_markdown(normalized, templates.get(template_name, {}), config)
    return {
        "template": template_name,
        "warnings": warnings,
        "preview": preview,
        "safe_user_output": preview,
    }


def compiler_render(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    template_name = optional_text(params, "template", str(config.get("default_template") or "blank"))
    target_format = optional_text(params, "format", str(config.get("default_format") or "md")).casefold().lstrip(".")
    if target_format not in SUPPORTED_FORMATS:
        raise ToolInputError(f"Unsupported report format: {target_format}")

    templates = load_templates(config)
    template = templates.get(template_name, {})
    normalized, errors, warnings = validate_content(params.get("content"), template)
    if errors:
        raise ToolInputError("; ".join(errors))

    output_path = report_output_path(config, target_format, normalized, optional_text(params, "output_path"))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    render_note = ""
    if bool(config.get("dry_run", False)):
        return {
            "rendered": False,
            "dry_run": True,
            "format": target_format,
            "template": template_name,
            "output_path": str(output_path),
            "warnings": warnings,
        }

    if target_format == "md":
        output_path.write_text(render_to_markdown(normalized, template, config), encoding="utf-8")
    elif target_format == "txt":
        output_path.write_text(render_to_text(normalized, template, config), encoding="utf-8")
    elif target_format == "docx":
        render_to_docx(normalized, template, config, output_path)
    elif target_format == "pdf":
        render_note = render_to_pdf(normalized, template, config, output_path)

    message = f"Rendered {normalized.get('title', 'report')} to {output_path}."
    if render_note:
        message = f"{message} {render_note}"
    return {
        "rendered": True,
        "format": target_format,
        "template": template_name,
        "output_path": str(output_path),
        "warnings": warnings,
        "safe_user_output": message,
    }


def config_path() -> Path:
    value = os.environ.get("JARVIS_REPORT_COMPILER_CONFIG", "").strip()
    return Path(value) if value else DEFAULT_CONFIG_PATH


def load_config() -> dict[str, Any]:
    path = config_path()
    if path.exists():
        data = json.loads(path.read_text(encoding="utf-8-sig"))
        if isinstance(data, dict):
            return ensure_config(data)
    data = ensure_config({})
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return data


def ensure_config(config: dict[str, Any]) -> dict[str, Any]:
    config.setdefault("agent_name", "Codex Report Compiler")
    config.setdefault("output_dir", "media/reports/output")
    config.setdefault("templates_dir", "templates")
    config.setdefault("default_format", "md")
    config.setdefault("default_template", "research_briefing")
    config.setdefault("docx_font", "Calibri")
    config.setdefault("docx_font_size", 11)
    config.setdefault("docx_margin_inches", 1.0)
    config.setdefault("include_toc", True)
    config.setdefault("include_page_numbers", True)
    config.setdefault("citation_style", "inline")
    config.setdefault("dry_run", False)
    return config


def dependency_status() -> dict[str, Any]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    return {
        "python_docx": find_spec("docx") is not None,
        "reportlab": find_spec("reportlab") is not None,
        "libreoffice": bool(soffice),
        "libreoffice_path": soffice or "",
        "markdown": find_spec("markdown") is not None,
    }


def load_templates(config: dict[str, Any]) -> dict[str, dict[str, Any]]:
    templates_dir = resolve_path(str(config.get("templates_dir") or "templates"))
    templates: dict[str, dict[str, Any]] = {}
    if not templates_dir.exists():
        return templates
    for path in sorted(templates_dir.glob("*.json"), key=lambda item: item.name.lower()):
        try:
            data = json.loads(path.read_text(encoding="utf-8-sig"))
        except json.JSONDecodeError:
            continue
        if isinstance(data, dict):
            name = str(data.get("name") or path.stem).strip()
            if name:
                templates[name] = data
    return templates


def validate_content(content: Any, template: dict[str, Any]) -> tuple[dict[str, Any], list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []
    if not isinstance(content, dict):
        return empty_content(), ["content must be an object"], warnings

    normalized = {
        "title": str(content.get("title") or "Untitled Report"),
        "subtitle": str(content.get("subtitle") or ""),
        "metadata": object_value(content.get("metadata")),
        "sections": [],
        "bibliography": [],
        "appendix": [],
    }
    if not str(content.get("title") or "").strip():
        warnings.append("title is missing; using Untitled Report")

    sections = content.get("sections")
    if sections is None:
        sections = []
    if not isinstance(sections, list):
        errors.append("sections must be a list")
        sections = []
    for index, raw_section in enumerate(sections, start=1):
        if not isinstance(raw_section, dict):
            warnings.append(f"section {index} ignored because it is not an object")
            continue
        normalized["sections"].append(normalize_section(raw_section, index))

    bibliography = content.get("bibliography")
    if bibliography is None:
        bibliography = []
    if not isinstance(bibliography, list):
        errors.append("bibliography must be a list")
        bibliography = []
    for index, raw_source in enumerate(bibliography, start=1):
        if isinstance(raw_source, dict):
            normalized["bibliography"].append(normalize_bibliography(raw_source, index))

    appendix = content.get("appendix")
    if isinstance(appendix, list):
        normalized["appendix"] = appendix
    elif appendix:
        warnings.append("appendix ignored because it is not a list")

    required_fields = template.get("required_fields", [])
    if isinstance(required_fields, list):
        for field in required_fields:
            if isinstance(field, str) and not normalized.get(field):
                warnings.append(f"template expects {field}")
    return normalized, errors, warnings


def empty_content() -> dict[str, Any]:
    return {"title": "Untitled Report", "subtitle": "", "metadata": {}, "sections": [], "bibliography": [], "appendix": []}


def normalize_section(raw: dict[str, Any], index: int) -> dict[str, Any]:
    level = raw.get("level", 1)
    try:
        clean_level = int(level)
    except (TypeError, ValueError):
        clean_level = 1
    return {
        "heading": str(raw.get("heading") or f"Section {index}"),
        "level": max(1, min(clean_level, 6)),
        "body": str(raw.get("body") or ""),
        "items": list_value(raw.get("items")),
        "table": raw.get("table") if isinstance(raw.get("table"), (dict, list)) else None,
        "citations": list_value(raw.get("citations")),
    }


def normalize_bibliography(raw: dict[str, Any], index: int) -> dict[str, Any]:
    return {
        "index": raw.get("index", index),
        "title": str(raw.get("title") or raw.get("url") or f"Source {index}"),
        "url": str(raw.get("url") or ""),
        "publisher": str(raw.get("publisher") or ""),
        "date": str(raw.get("date") or raw.get("published_date") or ""),
    }


def render_to_markdown(content: dict[str, Any], template: dict[str, Any], config: dict[str, Any]) -> str:
    lines: list[str] = [f"# {content['title']}"]
    if content.get("subtitle"):
        lines.extend(["", str(content["subtitle"])])
    metadata = object_value(content.get("metadata"))
    metadata_lines = metadata_as_lines(metadata)
    if metadata_lines:
        lines.extend(["", *metadata_lines])
    for section in content.get("sections", []):
        level = int(section.get("level") or 1)
        lines.extend(["", f"{'#' * min(level + 1, 6)} {section.get('heading', '')}"])
        body = str(section.get("body") or "").strip()
        if body:
            lines.extend(["", body])
        for item in section.get("items", []):
            lines.append(f"- {format_item(item)}")
        table_text = markdown_table(section.get("table"))
        if table_text:
            lines.extend(["", table_text])
        citations = citation_text(section.get("citations"))
        if citations:
            lines.extend(["", citations])
    bibliography = content.get("bibliography", [])
    if bibliography:
        lines.extend(["", "## Sources"])
        for source in bibliography:
            date = str(source.get("date") or "")
            suffix = f" ({date})" if date else ""
            url = str(source.get("url") or "")
            publisher = str(source.get("publisher") or "")
            label = str(source.get("title") or url or "Source")
            if url:
                lines.append(f"- [{source.get('index')}] {label}{suffix} - {publisher} - {url}".strip())
            else:
                lines.append(f"- [{source.get('index')}] {label}{suffix}")
    appendix = content.get("appendix", [])
    if appendix:
        lines.extend(["", "## Appendix"])
        for item in appendix:
            lines.append(format_item(item))
    return "\n".join(line.rstrip() for line in lines).strip() + "\n"


def render_to_text(content: dict[str, Any], template: dict[str, Any], config: dict[str, Any]) -> str:
    markdown = render_to_markdown(content, template, config)
    lines = []
    for line in markdown.splitlines():
        clean = line.lstrip("#").strip()
        if clean.startswith("- "):
            clean = "* " + clean[2:]
        lines.append(clean)
    return "\n".join(lines).strip() + "\n"


def render_to_docx(content: dict[str, Any], template: dict[str, Any], config: dict[str, Any], output_path: Path) -> None:
    if find_spec("docx") is None:
        raise ToolInputError("python-docx is required for DOCX output. Install python-docx or choose md/txt.")
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    margin = float(config.get("docx_margin_inches") or 1.0)
    for section in document.sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)
    style = document.styles["Normal"]
    style.font.name = str(config.get("docx_font") or "Calibri")
    style.font.size = Pt(float(config.get("docx_font_size") or 11))

    document.add_heading(str(content.get("title") or "Untitled Report"), 0)
    if content.get("subtitle"):
        document.add_paragraph(str(content.get("subtitle")))
    for line in metadata_as_lines(object_value(content.get("metadata"))):
        document.add_paragraph(line)
    for section in content.get("sections", []):
        document.add_heading(str(section.get("heading") or "Section"), level=min(int(section.get("level") or 1), 4))
        add_docx_body(document, str(section.get("body") or ""))
        for item in section.get("items", []):
            document.add_paragraph(format_item(item), style="List Bullet")
        add_docx_table(document, section.get("table"))
        citations = citation_text(section.get("citations"))
        if citations:
            document.add_paragraph(citations)
    if content.get("bibliography"):
        document.add_heading("Sources", level=1)
        for source in content.get("bibliography", []):
            document.add_paragraph(source_line(source), style="List Bullet")
    if content.get("appendix"):
        document.add_heading("Appendix", level=1)
        for item in content.get("appendix", []):
            document.add_paragraph(format_item(item))
    document.save(str(output_path))


def render_to_pdf(content: dict[str, Any], template: dict[str, Any], config: dict[str, Any], output_path: Path) -> str:
    if find_spec("reportlab") is not None:
        render_pdf_reportlab(content, template, config, output_path)
        return "Used reportlab."

    deps = dependency_status()
    if find_spec("docx") is None or not deps.get("libreoffice"):
        raise ToolInputError("PDF output requires reportlab, or python-docx plus LibreOffice/soffice.")

    with tempfile.TemporaryDirectory() as tmp:
        tmp_docx = Path(tmp) / "report.docx"
        render_to_docx(content, template, config, tmp_docx)
        command = [
            str(deps["libreoffice_path"]),
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_path.parent),
            str(tmp_docx),
        ]
        completed = subprocess.run(command, capture_output=True, text=True, timeout=120)
        converted = output_path.parent / "report.pdf"
        if completed.returncode != 0 or not converted.exists():
            raise ToolInputError(f"LibreOffice PDF conversion failed: {completed.stderr or completed.stdout}")
        if converted != output_path:
            if output_path.exists():
                output_path.unlink()
            converted.replace(output_path)
    return "Used LibreOffice conversion."


def render_pdf_reportlab(content: dict[str, Any], template: dict[str, Any], config: dict[str, Any], output_path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet

    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story: list[Any] = [Paragraph(escape(str(content.get("title") or "Untitled Report")), styles["Title"])]
    if content.get("subtitle"):
        story.extend([Spacer(1, 8), Paragraph(escape(str(content["subtitle"])), styles["Normal"])])
    for line in metadata_as_lines(object_value(content.get("metadata"))):
        story.append(Paragraph(escape(line), styles["Normal"]))
    for section in content.get("sections", []):
        story.extend([Spacer(1, 12), Paragraph(escape(str(section.get("heading") or "Section")), styles["Heading1"])])
        for paragraph in body_paragraphs(str(section.get("body") or "")):
            story.append(Paragraph(escape(paragraph), styles["BodyText"]))
        for item in section.get("items", []):
            story.append(Paragraph("- " + escape(format_item(item)), styles["BodyText"]))
        table_data = table_rows(section.get("table"))
        if table_data:
            table = Table(table_data)
            table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey)]))
            story.append(table)
        citations = citation_text(section.get("citations"))
        if citations:
            story.append(Paragraph(escape(citations), styles["BodyText"]))
    if content.get("bibliography"):
        story.extend([Spacer(1, 12), Paragraph("Sources", styles["Heading1"])])
        for source in content.get("bibliography", []):
            story.append(Paragraph(escape(source_line(source)), styles["BodyText"]))
    doc.build(story)


def add_docx_body(document: Any, body: str) -> None:
    for paragraph in body_paragraphs(body):
        document.add_paragraph(paragraph)


def add_docx_table(document: Any, table: Any) -> None:
    rows = table_rows(table)
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    docx_table = document.add_table(rows=len(rows), cols=column_count)
    docx_table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            docx_table.cell(row_index, column_index).text = str(row[column_index]) if column_index < len(row) else ""


def body_paragraphs(body: str) -> list[str]:
    paragraphs = []
    current: list[str] = []
    for line in body.splitlines():
        clean = line.strip()
        if not clean:
            if current:
                paragraphs.append(" ".join(current))
                current = []
            continue
        current.append(clean)
    if current:
        paragraphs.append(" ".join(current))
    return paragraphs


def markdown_table(table: Any) -> str:
    rows = table_rows(table)
    if not rows:
        return ""
    header = rows[0]
    output = ["| " + " | ".join(str(item) for item in header) + " |"]
    output.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows[1:]:
        output.append("| " + " | ".join(str(item) for item in row) + " |")
    return "\n".join(output)


def table_rows(table: Any) -> list[list[str]]:
    if isinstance(table, dict):
        headers = list_value(table.get("headers"))
        rows = table.get("rows")
        output = [[str(item) for item in headers]] if headers else []
        if isinstance(rows, list):
            for row in rows:
                if isinstance(row, list):
                    output.append([str(item) for item in row])
                elif isinstance(row, dict):
                    keys = headers or list(row.keys())
                    output.append([str(row.get(key, "")) for key in keys])
        return output
    if isinstance(table, list):
        return [[str(cell) for cell in row] for row in table if isinstance(row, list)]
    return []


def citation_text(citations: Any) -> str:
    values = list_value(citations)
    if not values:
        return ""
    return "Citations: " + ", ".join(format_item(item) for item in values)


def metadata_as_lines(metadata: dict[str, Any]) -> list[str]:
    lines = []
    for key, value in metadata.items():
        if value is None or value == "":
            continue
        title = str(key).replace("_", " ").title()
        lines.append(f"{title}: {value}")
    return lines


def source_line(source: dict[str, Any]) -> str:
    label = str(source.get("title") or source.get("url") or "Source")
    date = str(source.get("date") or "")
    publisher = str(source.get("publisher") or "")
    url = str(source.get("url") or "")
    parts = [f"[{source.get('index')}] {label}"]
    if date:
        parts.append(date)
    if publisher:
        parts.append(publisher)
    if url:
        parts.append(url)
    return " - ".join(parts)


def report_output_path(config: dict[str, Any], target_format: str, content: dict[str, Any], output_path: str) -> Path:
    if output_path:
        path = Path(output_path)
        if not path.is_absolute():
            path = resolve_path(output_path)
        if path.suffix.casefold().lstrip(".") != target_format:
            path = path.with_suffix("." + target_format)
        return path
    output_dir = resolve_path(str(config.get("output_dir") or "media/reports/output"))
    title = slug_text(str(content.get("title") or "report"))
    return output_dir / f"{timestamp_slug()}-{title}.{target_format}"


def slug_text(text: str) -> str:
    output = []
    last_dash = False
    for char in text.casefold():
        if char.isalnum():
            output.append(char)
            last_dash = False
        elif not last_dash:
            output.append("-")
            last_dash = True
    slug = "".join(output).strip("-")
    return slug or "report"


def timestamp_slug() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")


def resolve_path(path: str) -> Path:
    candidate = Path(path)
    return candidate if candidate.is_absolute() else (Path.cwd() / candidate).resolve()


def object_value(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def list_value(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def format_item(item: Any) -> str:
    if isinstance(item, dict):
        parts = []
        for key, value in item.items():
            if value is not None and value != "":
                parts.append(f"{key}: {value}")
        return "; ".join(parts)
    return str(item)
