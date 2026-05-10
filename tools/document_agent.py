from __future__ import annotations

import csv
import difflib
import hashlib
import json
import os
from datetime import datetime, timezone
from importlib.util import find_spec
from pathlib import Path
from typing import Any

from tools.path_resolver import resolve_local_path
from tools.registry import ToolInputError, optional_text, require_text


DEFAULT_CONFIG_PATH = Path("config/document_agent.json")


def document_status(_params: dict[str, Any] | None = None) -> dict[str, Any]:
    config = load_config()
    sessions = list_sessions(config)
    return {
        "summary": f"{config.get('agent_name', 'Document Agent')} ready.",
        "config_path": str(config_path()),
        "active_session_id": active_session_id(config),
        "session_count": len(sessions),
        "sessions": sessions,
        "supported_input_types": config.get("supported_input_types", []),
        "supported_output_types": config.get("supported_output_types", []),
        "dependencies": dependency_status(),
        "output_dir": str(resolve_path(str(config.get("output_dir")))),
        "session_dir": str(resolve_path(str(config.get("session_dir")))),
    }


def document_open(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    path = resolve_path(require_text(params, "path"))
    if not path.is_file():
        raise ToolInputError(f"path is not a file: {path}")
    source_type = path.suffix.casefold().lstrip(".")
    if source_type not in text_list(config.get("supported_input_types")):
        raise ToolInputError(f"Unsupported input type: {source_type}")

    loaded = load_document(path, config)
    session_id = optional_text(params, "session_id", "doc-" + short_hash(str(path.resolve()) + utc_now()))
    state = {
        "session_id": session_id,
        "source_path": str(path),
        "source_type": source_type,
        "opened_at": utc_now(),
        "structure": loaded["structure"],
        "metadata": loaded["metadata"],
        "content": loaded["content"],
        "pages": loaded["pages"],
        "tables": loaded["tables"],
        "images": loaded["images"],
        "sections": loaded["sections"],
        "annotations": [],
        "transforms_applied": [],
        "dirty": False,
        "output_path": None,
        "closed": False,
    }
    save_session(config, state)
    if params.get("make_active", True) is not False:
        set_active_session(config, session_id)
    return {
        "opened": True,
        "session_id": session_id,
        "source_path": str(path),
        "source_type": source_type,
        "structure": state["structure"],
        "summary": f"Opened {path.name} as session {session_id}.",
    }


def document_read(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    state = require_session(config, optional_text(params, "session_id"))
    max_chars = bounded_int(params.get("max_chars"), int(config.get("max_read_chars") or 12000), 100, 200000)
    section_heading = optional_text(params, "section_heading")
    if section_heading:
        content, bounds = section_content(state, section_heading)
    elif params.get("page_start") is not None or params.get("page_end") is not None:
        content, bounds = page_content(state, params)
    else:
        content, bounds = char_content(str(state.get("content") or ""), params)
    truncated = len(content) > max_chars
    return {
        "session_id": state["session_id"],
        "source_path": state.get("source_path"),
        "bounds": bounds,
        "chars": min(len(content), max_chars),
        "truncated": truncated,
        "content": content[:max_chars],
        "summary": f"Read {min(len(content), max_chars)} character(s) from session {state['session_id']}.",
    }


def document_extract(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    state = require_session(config, optional_text(params, "session_id"))
    operation = optional_text(params, "operation", "outline").casefold().replace(" ", "_")
    if operation in {"outline", "headings"}:
        return {
            "session_id": state["session_id"],
            "operation": operation,
            "headings": state.get("structure", {}).get("headings", []),
            "section_count": state.get("structure", {}).get("section_count", 0),
        }
    if operation == "tables":
        tables = state.get("tables", [])
        return {"session_id": state["session_id"], "operation": operation, "count": len(tables), "tables": tables}
    if operation == "images":
        images = state.get("images", [])
        return {"session_id": state["session_id"], "operation": operation, "count": len(images), "images": images}
    if operation == "citations":
        citations = extract_citations(str(state.get("content") or ""))
        return {"session_id": state["session_id"], "operation": operation, "count": len(citations), "citations": citations}
    if operation == "metadata":
        return {"session_id": state["session_id"], "operation": operation, "metadata": state.get("metadata", {})}
    if operation == "section":
        content, bounds = section_content(state, require_text(params, "section_heading"))
        return {
            "session_id": state["session_id"],
            "operation": operation,
            "bounds": bounds,
            "content": content[: bounded_int(params.get("max_chars"), int(config.get("max_read_chars") or 12000), 100, 200000)],
        }
    raise ToolInputError(f"Unsupported document extraction operation: {operation}")


def document_annotate(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    state = require_session(config, optional_text(params, "session_id"))
    note = require_text(params, "note")
    annotation = {
        "id": "ann-" + short_hash(state["session_id"] + note + utc_now()),
        "kind": optional_text(params, "kind", "comment"),
        "target": optional_text(params, "target"),
        "note": note,
        "location": params.get("location") if isinstance(params.get("location"), dict) else {},
        "text": optional_text(params, "text"),
        "created_at": utc_now(),
    }
    state.setdefault("annotations", []).append(annotation)
    state["dirty"] = True
    save_session(config, state)
    return {
        "annotated": True,
        "session_id": state["session_id"],
        "annotation": annotation,
        "annotation_count": len(state.get("annotations", [])),
    }


def document_transform(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    state = require_session(config, optional_text(params, "session_id"))
    operation = optional_text(params, "operation", "summarize").casefold().replace(" ", "_")
    source_text = transform_source_text(state, params, int(config.get("max_read_chars") or 12000))
    if operation in {"summarize", "summary"}:
        result_text = summarize_text(source_text, bounded_int(params.get("sentences"), 5, 1, 20))
    elif operation in {"executive_summary", "exec_summary"}:
        result_text = executive_summary_text(state, source_text)
    elif operation == "key_points":
        result_text = key_points_text(state, source_text, bounded_int(params.get("limit"), 8, 1, 30))
    elif operation in {"translate", "rewrite", "reading_level"}:
        result_text = model_followup_transform(operation, source_text, params)
    else:
        raise ToolInputError(f"Unsupported document transform operation: {operation}")

    transform = {
        "id": "tx-" + short_hash(state["session_id"] + operation + utc_now()),
        "operation": operation,
        "created_at": utc_now(),
        "parameters": public_params(params),
        "result": result_text,
    }
    state.setdefault("transforms_applied", []).append(transform)
    state["dirty"] = True
    save_session(config, state)
    return {
        "transformed": True,
        "session_id": state["session_id"],
        "transform": transform,
        "summary": f"Applied {operation} to session {state['session_id']}.",
    }


def document_compare(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    left = session_or_loaded_document(config, params.get("left_session_id"), params.get("left_path"))
    right = session_or_loaded_document(config, params.get("right_session_id"), params.get("right_path"))
    left_headings = heading_names(left)
    right_headings = heading_names(right)
    left_tokens = token_set(str(left.get("content") or ""))
    right_tokens = token_set(str(right.get("content") or ""))
    similarity = 1.0 if not left_tokens and not right_tokens else len(left_tokens & right_tokens) / max(1, len(left_tokens | right_tokens))
    diff_lines = list(
        difflib.unified_diff(
            str(left.get("content") or "").splitlines(),
            str(right.get("content") or "").splitlines(),
            fromfile=str(left.get("source_path") or "left"),
            tofile=str(right.get("source_path") or "right"),
            lineterm="",
        )
    )
    max_diff_chars = bounded_int(params.get("max_diff_chars"), 12000, 100, 50000)
    diff_text = "\n".join(diff_lines)
    return {
        "left": doc_label(left),
        "right": doc_label(right),
        "similarity_score": round(similarity, 3),
        "structural_diff": {
            "added_headings": [heading for heading in right_headings if heading not in left_headings],
            "removed_headings": [heading for heading in left_headings if heading not in right_headings],
            "shared_headings": [heading for heading in left_headings if heading in right_headings],
        },
        "content_diff": diff_text[:max_diff_chars],
        "truncated": len(diff_text) > max_diff_chars,
    }


def document_write(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    session_id = optional_text(params, "session_id")
    state: dict[str, Any] | None = None
    content = params.get("content")
    if not isinstance(content, dict):
        state = require_session(config, session_id)
        content = session_to_compiler_content(state, params)
    target_format = optional_text(params, "format", str(config.get("default_output_format") or "docx"))
    compiler_params = {
        "content": content,
        "format": target_format,
        "template": optional_text(params, "template", "blank"),
        "output_path": optional_text(params, "output_path"),
    }
    compiler_result = execute_registered_tool("compiler_render", compiler_params)
    if state is not None:
        state["output_path"] = compiler_result.get("output_path")
        state["dirty"] = False
        save_session(config, state)
    return {
        "written": True,
        "session_id": state.get("session_id") if state else "",
        "compiler_result": compiler_result,
        "output_path": compiler_result.get("output_path"),
        "format": compiler_result.get("format"),
        "safe_user_output": compiler_result.get("safe_user_output", ""),
    }


def document_merge(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    items: list[dict[str, Any]] = []
    for session_id in text_list(params.get("session_ids")):
        items.append(require_session(config, session_id))
    for raw_path in text_list(params.get("paths")):
        path = resolve_path(raw_path)
        loaded = load_document(path, config)
        loaded["source_path"] = str(path)
        loaded["source_type"] = path.suffix.casefold().lstrip(".")
        items.append(loaded)
    if not items:
        raise ToolInputError("session_ids or paths are required")
    title = optional_text(params, "title", "Merged Document")
    content_parts: list[str] = []
    sections: list[dict[str, Any]] = []
    for item in items:
        label = doc_label(item)
        body = str(item.get("content") or "")
        content_parts.append(f"# {label}\n\n{body}")
        sections.append({"heading": label, "level": 1, "start_char": 0, "end_char": len(body)})
    state = {
        "session_id": "doc-" + short_hash(title + utc_now()),
        "source_path": "",
        "source_type": "merged",
        "opened_at": utc_now(),
        "structure": structure_from_content("\n\n".join(content_parts), sections, [], default_language(config)),
        "metadata": {"title": title, "merged_count": len(items)},
        "content": "\n\n".join(content_parts),
        "pages": [],
        "tables": [],
        "images": [],
        "sections": sections,
        "annotations": [],
        "transforms_applied": [],
        "dirty": True,
        "output_path": None,
        "closed": False,
    }
    save_session(config, state)
    set_active_session(config, state["session_id"])
    return {"merged": True, "session_id": state["session_id"], "source_count": len(items), "structure": state["structure"]}


def document_close(params: dict[str, Any]) -> dict[str, Any]:
    config = load_config()
    state = require_session(config, optional_text(params, "session_id"))
    write_result: dict[str, Any] = {}
    if bool(params.get("save", False)):
        write_result = document_write(
            {
                "session_id": state["session_id"],
                "format": optional_text(params, "format", str(config.get("default_output_format") or "docx")),
                "output_path": optional_text(params, "output_path"),
            }
        )
        state = require_session(config, state["session_id"])
    state["closed"] = True
    save_session(config, state)
    if active_session_id(config) == state["session_id"]:
        clear_active_session(config)
    if bool(params.get("delete_session", False)):
        session_path(config, state["session_id"]).unlink(missing_ok=True)
    if bool(params.get("clean_cache", False)):
        clean_cache_dir(config)
    return {"closed": True, "session_id": state["session_id"], "saved": bool(write_result), "write_result": write_result}


def config_path() -> Path:
    value = os.environ.get("JARVIS_DOCUMENT_AGENT_CONFIG", "").strip()
    return Path(value) if value else DEFAULT_CONFIG_PATH


def load_config() -> dict[str, Any]:
    path = config_path()
    if path.exists():
        data = json.loads(path.read_text(encoding="utf-8-sig"))
        if isinstance(data, dict):
            return ensure_config(data)
    data = ensure_config({})
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return data


def ensure_config(config: dict[str, Any]) -> dict[str, Any]:
    config.setdefault("agent_name", "Codex Document Agent")
    config.setdefault("output_dir", "media/documents/output")
    config.setdefault("session_dir", "media/documents/sessions")
    config.setdefault("cache_dir", "media/documents/cache")
    config.setdefault("max_sessions", 5)
    config.setdefault("max_read_chars", 12000)
    config.setdefault("default_output_format", "docx")
    config.setdefault("supported_input_types", ["pdf", "docx", "md", "txt", "csv"])
    config.setdefault("supported_output_types", ["docx", "md", "pdf", "txt"])
    config.setdefault("summarize_model", "")
    config.setdefault("default_language", "en")
    config.setdefault("dry_run", False)
    return config


def dependency_status() -> dict[str, Any]:
    return {
        "python_docx": find_spec("docx") is not None,
        "pypdf": find_spec("pypdf") is not None,
        "pdfplumber": find_spec("pdfplumber") is not None,
        "markdown": find_spec("markdown") is not None,
        "report_compiler": registered_tool_ready("compiler_render"),
    }


def load_document(path: Path, config: dict[str, Any]) -> dict[str, Any]:
    suffix = path.suffix.casefold().lstrip(".")
    if suffix in {"txt", "md"}:
        content = path.read_text(encoding="utf-8-sig", errors="replace")
        sections = markdown_sections(content) if suffix == "md" else []
        return loaded_payload(path, content, [], [], [], sections, config, {"name": path.name})
    if suffix == "csv":
        content, tables = read_csv_document(path)
        return loaded_payload(path, content, [], tables, [], [], config, {"name": path.name})
    if suffix == "docx":
        return read_docx_document(path, config)
    if suffix == "pdf":
        return read_pdf_document(path, config)
    raise ToolInputError(f"Unsupported input type: {suffix}")


def loaded_payload(
    path: Path,
    content: str,
    pages: list[str],
    tables: list[dict[str, Any]],
    images: list[dict[str, Any]],
    sections: list[dict[str, Any]],
    config: dict[str, Any],
    metadata: dict[str, Any],
) -> dict[str, Any]:
    return {
        "source_path": str(path),
        "source_type": path.suffix.casefold().lstrip("."),
        "metadata": metadata,
        "content": content,
        "pages": pages,
        "tables": tables,
        "images": images,
        "sections": sections,
        "structure": structure_from_content(content, sections, pages, default_language(config)),
    }


def read_csv_document(path: Path) -> tuple[str, list[dict[str, Any]]]:
    rows: list[list[str]] = []
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
        reader = csv.reader(handle)
        for row in reader:
            rows.append([str(cell) for cell in row])
    content = "\n".join(", ".join(row) for row in rows)
    table = {"index": 1, "headers": rows[0] if rows else [], "rows": rows[1:] if len(rows) > 1 else []}
    return content, [table] if rows else []


def read_docx_document(path: Path, config: dict[str, Any]) -> dict[str, Any]:
    if find_spec("docx") is None:
        raise ToolInputError("python-docx is required to open DOCX files.")
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    sections: list[dict[str, Any]] = []
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        start = len("\n".join(parts))
        parts.append(text)
        style_name = getattr(paragraph.style, "name", "")
        if str(style_name).casefold().startswith("heading"):
            level = heading_level_from_style(str(style_name))
            sections.append({"heading": text, "level": level, "start_char": start, "end_char": start + len(text)})
    tables = []
    for index, table in enumerate(document.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append({"index": index, "headers": rows[0] if rows else [], "rows": rows[1:] if len(rows) > 1 else []})
    images = [{"index": index, "description": "Embedded image"} for index, _shape in enumerate(document.inline_shapes, start=1)]
    content = "\n".join(parts)
    metadata = {"name": path.name}
    return loaded_payload(path, content, [], tables, images, sections, config, metadata)


def read_pdf_document(path: Path, config: dict[str, Any]) -> dict[str, Any]:
    if find_spec("pdfplumber") is not None:
        return read_pdf_with_pdfplumber(path, config)
    if find_spec("pypdf") is not None:
        return read_pdf_with_pypdf(path, config)
    raise ToolInputError("PDF reading requires pdfplumber or pypdf.")


def read_pdf_with_pdfplumber(path: Path, config: dict[str, Any]) -> dict[str, Any]:
    import pdfplumber

    pages: list[str] = []
    tables: list[dict[str, Any]] = []
    with pdfplumber.open(str(path)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            pages.append(page.extract_text() or "")
            for table_index, table in enumerate(page.extract_tables() or [], start=1):
                headers = table[0] if table else []
                rows = table[1:] if len(table) > 1 else []
                tables.append({"page": page_index, "index": table_index, "headers": headers, "rows": rows})
    return loaded_payload(path, "\n\n".join(pages), pages, tables, [], [], config, {"name": path.name})


def read_pdf_with_pypdf(path: Path, config: dict[str, Any]) -> dict[str, Any]:
    import pypdf

    pages: list[str] = []
    with path.open("rb") as handle:
        reader = pypdf.PdfReader(handle)
        for page in reader.pages:
            pages.append(page.extract_text() or "")
    return loaded_payload(path, "\n\n".join(pages), pages, [], [], [], config, {"name": path.name})


def markdown_sections(content: str) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    cursor = 0
    for line in content.splitlines(keepends=True):
        stripped = line.lstrip()
        if stripped.startswith("#"):
            marks = 0
            for char in stripped:
                if char == "#":
                    marks += 1
                else:
                    break
            title = stripped[marks:].strip()
            if title:
                if sections:
                    sections[-1]["end_char"] = cursor
                sections.append({"heading": title, "level": max(1, min(marks, 6)), "start_char": cursor, "end_char": len(content)})
        cursor += len(line)
    return sections


def structure_from_content(content: str, sections: list[dict[str, Any]], pages: list[str], language: str) -> dict[str, Any]:
    return {
        "headings": [{"heading": item.get("heading", ""), "level": item.get("level", 1)} for item in sections],
        "section_count": len(sections),
        "page_count": len(pages),
        "word_count": len(content.split()),
        "language": language,
    }


def section_content(state: dict[str, Any], section_heading: str) -> tuple[str, dict[str, Any]]:
    target = section_heading.casefold()
    for section in state.get("sections", []):
        heading = str(section.get("heading") or "")
        if heading.casefold() == target or target in heading.casefold():
            start = int(section.get("start_char") or 0)
            end = int(section.get("end_char") or len(str(state.get("content") or "")))
            return str(state.get("content") or "")[start:end], {"type": "section", "heading": heading, "start_char": start, "end_char": end}
    raise ToolInputError(f"Section not found: {section_heading}")


def page_content(state: dict[str, Any], params: dict[str, Any]) -> tuple[str, dict[str, Any]]:
    pages = [str(page) for page in state.get("pages", [])]
    if not pages:
        raise ToolInputError("This session does not have page-level text.")
    start = bounded_int(params.get("page_start"), 1, 1, len(pages))
    end = bounded_int(params.get("page_end"), start, start, len(pages))
    selected = pages[start - 1 : end]
    return "\n\n".join(selected), {"type": "pages", "page_start": start, "page_end": end}


def char_content(content: str, params: dict[str, Any]) -> tuple[str, dict[str, Any]]:
    start = bounded_int(params.get("start_char"), 0, 0, len(content))
    end = bounded_int(params.get("end_char"), len(content), start, len(content))
    return content[start:end], {"type": "chars", "start_char": start, "end_char": end}


def extract_citations(content: str) -> list[dict[str, Any]]:
    citations = []
    seen: set[str] = set()
    for token in content.replace("\n", " ").split(" "):
        clean = token.strip().strip(".,;)]}")
        if not clean or clean in seen:
            continue
        if "://" in clean or clean.casefold().startswith("doi:"):
            seen.add(clean)
            citations.append({"text": clean})
    return citations


def transform_source_text(state: dict[str, Any], params: dict[str, Any], limit: int) -> str:
    if optional_text(params, "section_heading"):
        text, _bounds = section_content(state, optional_text(params, "section_heading"))
    else:
        text = str(state.get("content") or "")
    return text[: bounded_int(params.get("max_chars"), limit, 100, 200000)]


def summarize_text(text: str, sentence_count: int) -> str:
    sentences = split_sentences(text)
    if not sentences:
        return ""
    return " ".join(sentences[:sentence_count])


def executive_summary_text(state: dict[str, Any], text: str) -> str:
    title = Path(str(state.get("source_path") or "document")).name
    summary = summarize_text(text, 4)
    points = key_points_text(state, text, 5)
    return f"{title}\n\nSummary:\n{summary}\n\nKey points:\n{points}".strip()


def key_points_text(state: dict[str, Any], text: str, limit: int) -> str:
    headings = heading_names(state)
    if headings:
        return "\n".join(f"- {heading}" for heading in headings[:limit])
    sentences = split_sentences(text)
    return "\n".join(f"- {sentence}" for sentence in sentences[:limit])


def model_followup_transform(operation: str, text: str, params: dict[str, Any]) -> str:
    instruction = optional_text(params, "instruction", operation.replace("_", " "))
    return "\n".join(
        [
            f"Requested transform: {instruction}",
            "Use the assistant model to complete this transformation from the source excerpt below.",
            "",
            text,
        ]
    ).strip()


def split_sentences(text: str) -> list[str]:
    sentences: list[str] = []
    current: list[str] = []
    for char in text:
        current.append(char)
        if char in ".?!\n":
            sentence = " ".join("".join(current).split())
            if sentence:
                sentences.append(sentence)
            current = []
    tail = " ".join("".join(current).split())
    if tail:
        sentences.append(tail)
    return sentences


def session_to_compiler_content(state: dict[str, Any], params: dict[str, Any]) -> dict[str, Any]:
    title = optional_text(params, "title", Path(str(state.get("source_path") or "Document")).stem or "Document")
    sections = []
    document_sections = state.get("sections", [])
    if document_sections:
        for section in document_sections:
            start = int(section.get("start_char") or 0)
            end = int(section.get("end_char") or len(str(state.get("content") or "")))
            sections.append(
                {
                    "heading": str(section.get("heading") or "Section"),
                    "level": int(section.get("level") or 1),
                    "body": str(state.get("content") or "")[start:end],
                    "items": [],
                    "table": None,
                    "citations": [],
                }
            )
    else:
        sections.append({"heading": "Content", "level": 1, "body": str(state.get("content") or ""), "items": [], "table": None, "citations": []})
    for transform in state.get("transforms_applied", []):
        if isinstance(transform, dict):
            sections.append(
                {
                    "heading": str(transform.get("operation") or "Transform"),
                    "level": 1,
                    "body": str(transform.get("result") or ""),
                    "items": [],
                    "table": None,
                    "citations": [],
                }
            )
    return {
        "title": title,
        "subtitle": optional_text(params, "subtitle"),
        "metadata": {
            "source_path": state.get("source_path", ""),
            "session_id": state.get("session_id", ""),
            "date": utc_now(),
        },
        "sections": sections,
        "bibliography": [],
        "appendix": state.get("annotations", []),
    }


def execute_registered_tool(name: str, params: dict[str, Any]) -> dict[str, Any]:
    from tools.registry import discover_tools

    payload = json.loads(discover_tools().execute(name, params))
    if not payload.get("ok"):
        raise ToolInputError(str(payload.get("error") or f"{name} failed"))
    result = payload.get("result")
    if not isinstance(result, dict):
        raise ToolInputError(f"{name} returned an invalid result")
    return result


def registered_tool_ready(name: str) -> bool:
    try:
        from tools.registry import discover_tools

        return name in [tool.name for tool in discover_tools().visible_tools()]
    except Exception:
        return False


def session_or_loaded_document(config: dict[str, Any], session_value: Any, path_value: Any) -> dict[str, Any]:
    session_id = str(session_value or "").strip()
    if session_id:
        return require_session(config, session_id)
    path_text = str(path_value or "").strip()
    if path_text:
        path = resolve_path(path_text)
        loaded = load_document(path, config)
        loaded["source_path"] = str(path)
        return loaded
    raise ToolInputError("session id or path is required for comparison")


def require_session(config: dict[str, Any], session_id: str = "") -> dict[str, Any]:
    clean = session_id.strip() if session_id else active_session_id(config)
    if not clean:
        raise ToolInputError("No document session is active; open a document first or pass session_id.")
    path = session_path(config, clean)
    if not path.exists():
        raise ToolInputError(f"Document session not found: {clean}")
    data = json.loads(path.read_text(encoding="utf-8-sig"))
    if not isinstance(data, dict):
        raise ToolInputError(f"Document session is not valid JSON: {clean}")
    return data


def save_session(config: dict[str, Any], state: dict[str, Any]) -> None:
    path = session_path(config, str(state.get("session_id") or ""))
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")
    prune_sessions(config)


def prune_sessions(config: dict[str, Any]) -> None:
    maximum = bounded_int(config.get("max_sessions"), 5, 1, 50)
    files = sorted(session_dir(config).glob("doc-*.json"), key=lambda path: path.stat().st_mtime, reverse=True)
    for path in files[maximum:]:
        if path.stem != active_session_id(config):
            path.unlink(missing_ok=True)


def list_sessions(config: dict[str, Any]) -> list[dict[str, Any]]:
    sessions = []
    for path in sorted(session_dir(config).glob("doc-*.json"), key=lambda item: item.stat().st_mtime, reverse=True):
        try:
            data = json.loads(path.read_text(encoding="utf-8-sig"))
        except json.JSONDecodeError:
            continue
        if isinstance(data, dict):
            sessions.append(
                {
                    "session_id": data.get("session_id", ""),
                    "source_path": data.get("source_path", ""),
                    "source_type": data.get("source_type", ""),
                    "opened_at": data.get("opened_at", ""),
                    "dirty": data.get("dirty", False),
                    "closed": data.get("closed", False),
                    "word_count": data.get("structure", {}).get("word_count", 0),
                }
            )
    return sessions


def active_session_id(config: dict[str, Any]) -> str:
    path = active_session_path(config)
    if not path.exists():
        return ""
    try:
        data = json.loads(path.read_text(encoding="utf-8-sig"))
    except json.JSONDecodeError:
        return ""
    return str(data.get("session_id") or "") if isinstance(data, dict) else ""


def set_active_session(config: dict[str, Any], session_id: str) -> None:
    path = active_session_path(config)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps({"session_id": session_id}, indent=2), encoding="utf-8")


def clear_active_session(config: dict[str, Any]) -> None:
    active_session_path(config).unlink(missing_ok=True)


def session_dir(config: dict[str, Any]) -> Path:
    return resolve_path(str(config.get("session_dir") or "media/documents/sessions"))


def session_path(config: dict[str, Any], session_id: str) -> Path:
    clean = session_id.strip()
    if not clean:
        raise ToolInputError("session_id is required")
    return session_dir(config) / f"{clean}.json"


def active_session_path(config: dict[str, Any]) -> Path:
    return session_dir(config) / "_active.json"


def clean_cache_dir(config: dict[str, Any]) -> None:
    cache_dir = resolve_path(str(config.get("cache_dir") or "media/documents/cache"))
    if not cache_dir.exists():
        return
    for path in cache_dir.iterdir():
        if path.is_file():
            path.unlink(missing_ok=True)


def resolve_path(value: str) -> Path:
    return resolve_local_path(value)


def heading_level_from_style(style_name: str) -> int:
    for part in style_name.split():
        try:
            return max(1, min(int(part), 6))
        except ValueError:
            continue
    return 1


def heading_names(state: dict[str, Any]) -> list[str]:
    headings = state.get("structure", {}).get("headings", [])
    return [str(item.get("heading") or "") for item in headings if isinstance(item, dict) and str(item.get("heading") or "")]


def token_set(text: str) -> set[str]:
    tokens: set[str] = set()
    for raw in text.casefold().split():
        clean = "".join(char for char in raw if char.isalnum())
        if clean:
            tokens.add(clean)
    return tokens


def doc_label(state: dict[str, Any]) -> str:
    path = str(state.get("source_path") or "")
    if path:
        return Path(path).name
    return str(state.get("session_id") or state.get("metadata", {}).get("title") or "document")


def public_params(params: dict[str, Any]) -> dict[str, Any]:
    return {key: value for key, value in params.items() if key != "content"}


def default_language(config: dict[str, Any]) -> str:
    return str(config.get("default_language") or "en")


def text_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def bounded_int(value: Any, fallback: int, minimum: int, maximum: int) -> int:
    if isinstance(value, int):
        number = value
    elif isinstance(value, str) and value.strip():
        try:
            number = int(value.strip())
        except ValueError:
            number = fallback
    else:
        number = fallback
    return max(minimum, min(maximum, number))


def short_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:12]


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()
